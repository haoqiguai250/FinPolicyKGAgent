"""
材料自动生成模块

根据政策要求 + 企业画像，LLM 生成文档内容，
输出可下载的 PDF/Word 文件。
"""

import json
import time
import uuid
from pathlib import Path
from typing import Optional

from loguru import logger

from config.settings import settings


class MaterialDocumentGenerator:
    """材料文档生成器"""

    def __init__(self):
        self.output_dir = settings.MATERIALS_OUTPUT_DIR
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_document(
        self,
        opportunity: dict,
        materials: list[dict],
        enterprise_profile: dict,
        doc_type: str = "docx",
    ) -> list[dict]:
        """
        完整流水线：为每条材料生成文档

        Args:
            opportunity: 申报机会（DB 行）
            materials: 材料清单（DB 行列表）
            enterprise_profile: 企业画像
            doc_type: 输出格式 docx / pdf

        Returns:
            生成的文档记录列表（用于写 DB）
        """
        from src.extraction.llm_client import get_llm_client
        llm = get_llm_client()

        results = []
        opp_id = opportunity["opportunity_id"]
        policy_name = opportunity["policy_name"]

        for mat in materials:
            mat_name = mat["material_name"]
            mat_id = mat.get("material_id", "")
            logger.info(f"生成文档: {mat_name} (opportunity={opp_id})")

            try:
                # 1. LLM 生成文档内容
                content = self._generate_content(
                    llm=llm,
                    policy_name=policy_name,
                    material_name=mat_name,
                    enterprise_profile=profile_simplified(enterprise_profile),
                    opportunity=opportunity,
                )

                # 2. 生成文件
                ts = int(time.time())
                safe_name = _safe_filename(mat_name)
                filename = f"{opp_id[:8]}_{safe_name}_{ts}.{doc_type}"
                file_path = self.output_dir / filename

                if doc_type == "docx":
                    self._generate_docx(file_path, mat_name, policy_name, content, enterprise_profile)
                else:
                    self._generate_pdf(file_path, mat_name, policy_name, content, enterprise_profile)

                file_size = file_path.stat().st_size

                results.append({
                    "doc_id": str(uuid.uuid4()),
                    "opportunity_id": opp_id,
                    "material_id": mat_id,
                    "doc_name": filename,
                    "doc_type": doc_type,
                    "file_path": str(file_path),
                    "file_size": file_size,
                    "status": "generated",
                })
                logger.info(f"文档已生成: {filename} ({file_size} bytes)")

            except Exception as e:
                logger.error(f"生成文档失败 [{mat_name}]: {e}")
                results.append({
                    "doc_id": str(uuid.uuid4()),
                    "opportunity_id": opp_id,
                    "material_id": mat_id,
                    "doc_name": f"{mat_name}.{doc_type}",
                    "doc_type": doc_type,
                    "file_path": "",
                    "file_size": 0,
                    "status": "error",
                })

        return results

    def _generate_content(
        self,
        llm,
        policy_name: str,
        material_name: str,
        enterprise_profile: dict,
        opportunity: dict,
    ) -> dict:
        """LLM 生成结构化文档内容"""
        prompt = f"""请根据以下信息，生成"{material_name}"的文档内容。

政策名称: {policy_name}
企业信息: {json.dumps(enterprise_profile, ensure_ascii=False)}
申报说明: {opportunity.get('match_explanation', '')}
建议: {opportunity.get('suggestions', '')}

请生成该材料的结构化内容，严格按 JSON 格式输出：
{{
  "title": "文档标题",
  "sections": [
    {{"heading": "章节标题", "content": "章节内容（支持多行文本）"}}
  ],
  "table_data": [
    {{"field": "字段名", "value": "对应值"}}
  ],
  "notes": "补充说明"
}}

要求：
- 根据企业画像自动填充可确定的字段
- 无法确定的字段用"[待填写]"标注
- 内容应符合政府申报材料的规范格式
- 表格数据包含企业基本信息（名称、统一社会信用代码、注册地址等）"""

        result = llm.chat_json(
            system_prompt="你是一个政府政策申报材料撰写专家，熟悉各类申报材料的格式和要求。",
            user_prompt=prompt,
            temperature=0.2,
        )

        if isinstance(result, dict):
            return result
        return {"title": material_name, "sections": [], "table_data": [], "notes": ""}

    def _generate_docx(
        self,
        file_path: Path,
        material_name: str,
        policy_name: str,
        content: dict,
        enterprise_profile: dict,
    ) -> None:
        """用 python-docx 生成 Word 文件"""
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 页面设置
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

        # 标题
        title = content.get("title", material_name)
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 政策信息
        doc.add_paragraph(f"对应政策：{policy_name}", style="Normal")

        # 企业信息表格
        table_data = content.get("table_data", [])
        if table_data:
            doc.add_heading("企业基本信息", level=2)
            table = doc.add_table(rows=len(table_data), cols=2, style="Table Grid")
            for i, row_data in enumerate(table_data):
                table.rows[i].cells[0].text = str(row_data.get("field", ""))
                table.rows[i].cells[1].text = str(row_data.get("value", ""))

        # 章节内容
        sections = content.get("sections", [])
        for sec in sections:
            heading_text = sec.get("heading", "")
            if heading_text:
                doc.add_heading(heading_text, level=2)
            content_text = sec.get("content", "")
            if content_text:
                for line in content_text.split("\n"):
                    doc.add_paragraph(line.strip())

        # 补充说明
        notes = content.get("notes", "")
        if notes:
            doc.add_heading("补充说明", level=2)
            doc.add_paragraph(notes)

        doc.save(str(file_path))

    def _generate_pdf(
        self,
        file_path: Path,
        material_name: str,
        policy_name: str,
        content: dict,
        enterprise_profile: dict,
    ) -> None:
        """用 reportlab 生成 PDF 文件"""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        # 尝试注册中文字体
        try:
            pdfmetrics.registerFont(TTFont("SimSun", "simsun.ttc"))
            cn_font = "SimSun"
        except Exception:
            cn_font = "Helvetica"

        doc = SimpleDocTemplate(
            str(file_path),
            pagesize=A4,
            leftMargin=3 * cm,
            rightMargin=3 * cm,
            topMargin=2.5 * cm,
            bottomMargin=2.5 * cm,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "CNTitle",
            parent=styles["Title"],
            fontName=cn_font,
            fontSize=16,
            alignment=1,
        )
        heading_style = ParagraphStyle(
            "CNHeading",
            parent=styles["Heading2"],
            fontName=cn_font,
            fontSize=13,
        )
        normal_style = ParagraphStyle(
            "CNNormal",
            parent=styles["Normal"],
            fontName=cn_font,
            fontSize=10.5,
            leading=18,
        )

        story = []

        # 标题
        title = content.get("title", material_name)
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph(f"对应政策：{policy_name}", normal_style))
        story.append(Spacer(1, 0.5 * cm))

        # 企业信息表格
        table_data = content.get("table_data", [])
        if table_data:
            story.append(Paragraph("企业基本信息", heading_style))
            story.append(Spacer(1, 0.3 * cm))
            pdf_table_data = [["字段", "值"]]
            for row in table_data:
                pdf_table_data.append([
                    str(row.get("field", "")),
                    str(row.get("value", "")),
                ])
            t = Table(pdf_table_data, colWidths=[5 * cm, 10 * cm])
            t.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTNAME", (0, 0), (-1, -1), cn_font),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ]))
            story.append(t)
            story.append(Spacer(1, 0.5 * cm))

        # 章节内容
        sections = content.get("sections", [])
        for sec in sections:
            heading_text = sec.get("heading", "")
            if heading_text:
                story.append(Paragraph(heading_text, heading_style))
                story.append(Spacer(1, 0.2 * cm))
            content_text = sec.get("content", "")
            if content_text:
                for line in content_text.split("\n"):
                    line = line.strip()
                    if line:
                        story.append(Paragraph(line, normal_style))
                story.append(Spacer(1, 0.3 * cm))

        # 补充说明
        notes = content.get("notes", "")
        if notes:
            story.append(Paragraph("补充说明", heading_style))
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph(notes, normal_style))

        doc.build(story)


def profile_simplified(profile: dict) -> dict:
    """精简企业画像，去掉空值字段"""
    return {k: v for k, v in profile.items() if v and v != "" and v != []}


def _safe_filename(name: str) -> str:
    """将名称转为安全文件名"""
    import re
    safe = re.sub(r'[<>:"/\\|?*\s]', '_', name)
    return safe[:30] if safe else "document"
