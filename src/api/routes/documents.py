"""
文档生成路由 — 材料文档自动生成 + 下载 + 演示模式
"""

import os
from pathlib import Path

from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel

from src.core.logger import logger

router = APIRouter()


class DocumentGenerateRequest(BaseModel):
    material_ids: list[str] = []


class DemoDocRequest(BaseModel):
    policy_name: str = ""
    doc_type: str = "application"
    materials: list[str] = []
    steps: list[str] = []
    amount_detail: str = ""
    deadline: str = ""
    department: str = ""
    enterprise_name: str = ""
    enterprise_region: str = ""
    enterprise_type: str = ""
    enterprise_industry: str = ""
    enterprise_employees: str = ""
    enterprise_revenue: str = ""


def _get_db():
    from src.api.server import get_db
    db = get_db()
    if not db:
        raise HTTPException(status_code=503, detail="数据库未初始化")
    return db


@router.post("/opportunities/{opportunity_id}/documents/generate")
async def generate_documents(opportunity_id: str, req: DocumentGenerateRequest):
    """为申报机会自动生成材料文档"""
    db = _get_db()
    opp = db.get_opportunity(opportunity_id)
    if not opp:
        raise HTTPException(status_code=404, detail="申报机会不存在")

    materials = db.list_materials(opportunity_id)
    if not materials:
        raise HTTPException(status_code=400, detail="材料清单为空，请先生成材料清单")

    if req.material_ids:
        materials = [m for m in materials if m["material_id"] in req.material_ids]

    from src.decision.material_generator import MaterialDocumentGenerator
    generator = MaterialDocumentGenerator()

    enterprise = db.get_enterprise_profile(opp["enterprise_id"]) or {}

    results = generator.generate_document(
        opportunity=opp,
        materials=materials,
        enterprise_profile=enterprise,
        doc_type="docx",
    )

    saved = []
    for r in results:
        try:
            with db.get_conn() as conn:
                conn.execute(
                    """INSERT INTO generated_documents (doc_id, opportunity_id, material_id,
                    doc_name, doc_type, file_path, file_size, status, created_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, datetime('now','localtime'))""",
                    (r["doc_id"], opportunity_id, r.get("material_id", ""),
                     r["doc_name"], r["doc_type"], r.get("file_path", ""),
                     r.get("file_size", 0), r["status"]))
                conn.commit()
            saved.append(r)
        except Exception as e:
            logger.warning(f"保存文档记录失败: {e}")

    for mat in materials:
        try:
            db.update_material(mat["material_id"], status="ready")
        except Exception:
            pass

    return {
        "opportunity_id": opportunity_id,
        "generated_count": len(saved),
        "documents": saved,
    }


@router.get("/opportunities/{opportunity_id}/documents")
async def list_documents(opportunity_id: str):
    """获取已生成文档列表"""
    db = _get_db()
    docs = db.list_generated_documents(opportunity_id)
    return {
        "opportunity_id": opportunity_id,
        "total": len(docs),
        "documents": docs,
    }


@router.get("/documents/{doc_id}/download")
async def download_document(doc_id: str):
    """下载已生成的文档（防路径穿越）"""
    db = _get_db()
    doc = db.get_generated_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    file_path = doc.get("file_path", "")
    if not file_path:
        raise HTTPException(status_code=404, detail="文件路径为空")

    resolved = Path(file_path).resolve()
    from config.settings import settings
    output_dir = settings.MATERIALS_OUTPUT_DIR.resolve()
    if not str(resolved).startswith(str(output_dir)):
        raise HTTPException(status_code=403, detail="非法文件路径")

    if not resolved.exists():
        raise HTTPException(status_code=404, detail="文件不存在")

    if doc.get("doc_type") == "pdf":
        media_type = "application/pdf"
    else:
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return FileResponse(
        path=str(resolved),
        filename=doc["doc_name"],
        media_type=media_type,
    )


@router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    """删除已生成的文档"""
    db = _get_db()
    doc = db.get_generated_document(doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    file_path = doc.get("file_path", "")
    if file_path:
        try:
            resolved = Path(file_path).resolve()
            from config.settings import settings
            output_dir = settings.MATERIALS_OUTPUT_DIR.resolve()
            if str(resolved).startswith(str(output_dir)) and resolved.exists():
                resolved.unlink()
        except Exception as e:
            logger.warning(f"删除文件失败: {e}")

    with db.get_conn() as conn:
        conn.execute("DELETE FROM generated_documents WHERE doc_id = ?", (doc_id,))

    return {"status": "ok", "message": "文档已删除"}


@router.post("/demo/documents/generate")
async def demo_generate_document(req: DemoDocRequest):
    """演示模式：用 python-docx 生成真实 Word 文档，直接返回文件下载"""
    from docx import Document
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import time as _time
    from config.settings import settings

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    pname = req.policy_name or "政策申报"

    if req.doc_type == "application":
        title = doc.add_heading(pname[:30] + "\n申报书", level=1)
    else:
        title = doc.add_heading("承 诺 函", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if req.doc_type == "application":
        # ── 申报单位信息（使用画像数据） ──
        doc.add_heading("一、申报单位信息", level=2)
        info_table = doc.add_table(rows=7, cols=2, style="Table Grid")
        info = [
            ("单位名称", req.enterprise_name or "深圳智创科技有限公司"),
            ("统一社会信用代码", "91440300XXXXXXXXXX"),
            ("注册地址", req.enterprise_region or "深圳市南山区"),
            ("所属行业", req.enterprise_industry or "人工智能/信息技术"),
            ("企业类型", req.enterprise_type or "民营科技企业"),
            ("职工人数", (req.enterprise_employees or "380") + "人"),
            ("上年度营收", (req.enterprise_revenue or "12000") + "万元"),
        ]
        for i, (k, v) in enumerate(info):
            info_table.rows[i].cells[0].text = k
            info_table.rows[i].cells[1].text = v

        # ── 申报项目信息 ──
        doc.add_paragraph("")
        doc.add_heading("二、申报项目信息", level=2)
        proj_table = doc.add_table(rows=5, cols=2, style="Table Grid")
        proj = [
            ("政策名称", pname),
            ("主管部门", req.department or "待确认"),
            ("预计资助金额", req.amount_detail or "按政策核定"),
            ("申报截止日期", req.deadline or "暂无"),
            ("申报平台", "待确认"),
        ]
        for i, (k, v) in enumerate(proj):
            proj_table.rows[i].cells[0].text = k
            proj_table.rows[i].cells[1].text = v

        # ── 所需材料（从政策数据填入） ──
        if req.materials:
            doc.add_paragraph("")
            doc.add_heading("三、所需申报材料", level=2)
            for idx, mat in enumerate(req.materials, 1):
                doc.add_paragraph(f"{idx}. {mat}", style="List Number")
        else:
            doc.add_paragraph("")
            doc.add_heading("三、所需申报材料", level=2)
            doc.add_paragraph("（根据申报指南准备）")

        # ── 申报步骤 ──
        if req.steps:
            doc.add_paragraph("")
            doc.add_heading("四、申报流程", level=2)
            for idx, step in enumerate(req.steps, 1):
                doc.add_paragraph(f"{idx}. {step}", style="List Number")

        doc.add_paragraph("")
        doc.add_paragraph("以上信息真实有效，如有虚假愿承担法律责任。")
    else:
        # ── 承诺函 ──
        doc.add_paragraph("深圳市科技创新局：")
        doc.add_paragraph("    我单位（" + (req.enterprise_name or "深圳智创科技有限公司") + "）就申报 " + pname[:20] + " 作出以下承诺：")
        items = [
            "所提交材料真实完整，无弄虚作假",
            "近三年无重大安全环保事故",
            "项目经费专款专用",
            "违反承诺愿承担法律责任并退回资金",
        ]
        for i, item in enumerate(items, 1):
            doc.add_paragraph("    " + str(i) + ". " + item)
        doc.add_paragraph("")
        doc.add_paragraph("承诺单位（盖章）：" + (req.enterprise_name or "深圳智创科技有限公司"))
        doc.add_paragraph("法定代表人（签字）：")
        doc.add_paragraph("日期：" + _time.strftime("%Y年%m月%d日"))

    ts = int(_time.time())
    filename = "demo_" + req.doc_type + "_" + str(ts) + ".docx"
    output_dir = settings.MATERIALS_OUTPUT_DIR
    output_dir.mkdir(parents=True, exist_ok=True)
    file_path = output_dir / filename
    doc.save(str(file_path))
    logger.info(f"演示文档已生成: {filename} ({file_path.stat().st_size} bytes)")

    return FileResponse(
        path=str(file_path),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
